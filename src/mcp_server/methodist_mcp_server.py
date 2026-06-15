from pathlib import Path
from typing import Optional

from mcp.server.fastmcp import FastMCP

from core.document_environment import DocumentEnvironment, DocumentRequest
from core.sandbox import Sandbox
from drivers.com_driver import COMDriver
from drivers.native_driver import NativeDriver

mcp = FastMCP("methodist-agent")


def _project_root() -> Path:
    return Path.home() / "Documents" / "Методист-Агент"


def _get_env(output_path: Optional[str] = None) -> DocumentEnvironment:
    return DocumentEnvironment(
        sandbox=Sandbox(_project_root()),
        drivers=[COMDriver(), NativeDriver()],
    )


def _normalize_path(path: str) -> Path:
    target = Path(path)
    if not target.is_absolute():
        target = _project_root() / target
    return _get_env().sandbox.normalize(target)


@mcp.tool()
def list_documents(folder: Optional[str] = None) -> list[str]:
    """List document files in a workspace folder."""
    root = _normalize_path(folder) if folder else _project_root()
    if not root.exists():
        return []
    return [str(p) for p in root.rglob("*") if p.is_file()]


@mcp.tool()
def create_document(doc_type: str, output_path: str, parameters: dict) -> str:
    """Create a document from parameters."""
    env = _get_env(output_path)
    req = DocumentRequest(
        action="create",
        doc_type=doc_type,
        output_path=output_path,
        parameters=parameters,
    )
    result = env.execute(req)
    if result.success:
        return f"Created: {result.output_path}"
    return f"Failed: {result.message}"


@mcp.tool()
def read_document(path: str) -> str:
    """Read text from a DOCX file."""
    from docx import Document

    doc_path = _normalize_path(path)
    doc = Document(str(doc_path))
    return "\n".join(p.text for p in doc.paragraphs)


@mcp.tool()
def convert_to_pdf(input_path: str, output_path: str) -> str:
    """Convert a document to PDF (stub)."""
    return f"PDF conversion not yet implemented for {input_path}"


@mcp.resource("doc://{relative_path}")
def get_document_resource(relative_path: str) -> str:
    """Read-only access to a project document."""
    path = _normalize_path(relative_path)
    if not path.exists():
        return f"Document not found: {relative_path}"
    if path.suffix == ".docx":
        from docx import Document

        doc = Document(str(path))
        return "\n".join(p.text for p in doc.paragraphs)
    return path.read_text(encoding="utf-8", errors="ignore")


@mcp.prompt()
def generate_curriculum(subject: str, hours: int) -> str:
    """Prompt template for generating a curriculum."""
    return f"Создай рабочую программу дисциплины '{subject}' объёмом {hours} часов."


@mcp.prompt()
def adapt_to_fgos(source_path: str, fgos_version: str = "3++") -> str:
    """Prompt template for adapting a document to FGOS."""
    return f"Адаптируй документ {source_path} под требования ФГОС {fgos_version}."


if __name__ == "__main__":
    mcp.run(transport="stdio")
