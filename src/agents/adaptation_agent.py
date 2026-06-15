"""
AdaptationAgent — агент адаптации документов для методистов.

Поддерживает:
- Адаптацию документов под новые шаблоны
- Сравнение версий документов
- Массовую замену текста
- Обновление нормативных ссылок
- Применение правил форматирования

Режимы работы:
- native: python-docx (кроссплатформенный)
- com: pywin32 COM (Windows, полная поддержка форматирования)
- auto: COM если доступен, иначе native
"""

import difflib
import logging
from pathlib import Path
from typing import Dict, Optional, Any
from datetime import datetime

from core.config import Config, get_output_dir, get_templates_dir
from core.ui_text import error_file_not_found, error_generic, error_no_office_fallback

logger = logging.getLogger(__name__)

# Optional Windows COM support
try:
    import win32com.client as win32
    _HAS_WIN32 = True
except ImportError:
    _HAS_WIN32 = False
    logger.debug("pywin32 не установлен — COM-автоматизация недоступна")

# Native .docx support
try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    _HAS_DOCX = True
except ImportError:
    _HAS_DOCX = False
    logger.warning("python-docx не установлен — нативная обработка .docx недоступна")


class AdaptationError(Exception):
    """Ошибка адаптации документа."""
    pass


class AdaptationAgent:
    """Агент для адаптации и обработки методических документов."""

    def __init__(self, config: Config):
        self.config = config
        self.output_dir = get_output_dir(config)
        self.templates_dir = get_templates_dir(config)
        self.mode = self._resolve_mode()
        logger.info(f"AdaptationAgent инициализирован (режим: {self.mode})")

    def _resolve_mode(self) -> str:
        """Определить рабочий режим обработки документов."""
        preferred = self.config.documents.preferred_mode
        if preferred == "com" and not _HAS_WIN32:
            logger.warning(error_no_office_fallback())
            return "native" if _HAS_DOCX else "none"
        if preferred == "native" and not _HAS_DOCX:
            logger.warning("Запрошен native-режим, но python-docx не установлен.")
            return "none"
        if preferred == "auto":
            if _HAS_WIN32:
                return "com"
            if _HAS_DOCX:
                return "native"
            return "none"
        return preferred

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def execute(self, task_type: str, params: Dict[str, Any]) -> Any:
        """Диспетчер задач адаптации."""
        logger.info(f"Выполнение задачи: {task_type}")

        dispatch = {
            "adapt_document": self.adapt_document,
            "compare_documents": lambda p: self.compare_documents(
                p.get("old_path"), p.get("new_path")
            ),
            "batch_replace": lambda p: self.batch_replace(
                p.get("doc_path"), p.get("replacements", {})
            ),
            "update_references": lambda p: self.update_references(p.get("doc_path")),
            "apply_formatting_rules": lambda p: self.apply_formatting_rules(
                p.get("doc_path"), p.get("rules", {})
            ),
        }

        handler = dispatch.get(task_type)
        if not handler:
            raise AdaptationError(f"Неизвестный тип задачи: {task_type}")

        try:
            return handler(params)
        except Exception as e:
            logger.exception(f"Ошибка при выполнении {task_type}")
            if isinstance(e, AdaptationError):
                raise
            friendly_action = {
                "adapt_document": "адаптировать документ",
                "compare_documents": "сравнить документы",
                "batch_replace": "выполнить замену текста",
                "update_references": "обновить ссылки",
                "apply_formatting_rules": "применить правила форматирования",
            }.get(task_type, task_type)
            raise AdaptationError(error_generic(friendly_action, str(e))) from e

    def adapt_document(self, params: Dict[str, Any]) -> str:
        """
        Адаптировать документ под новые требования.

        Параметры:
            input_file: путь к исходному документу
            template:   путь к новому шаблону (опционально)
            output:     путь для сохранения результата (опционально)
            rules:      словарь правил адаптации (опционально)

        Возвращает путь к созданному файлу.
        """
        input_file = params.get("input_file")
        template = params.get("template")
        output = params.get("output")
        rules = params.get("rules", {})

        if not input_file:
            raise AdaptationError("Не указан input_file")

        input_path = Path(input_file)
        if not input_path.exists():
            raise AdaptationError(error_file_not_found(input_file))

        # Определить путь выходного файла
        if output:
            output_path = Path(output)
        else:
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            stem = input_path.stem
            output_path = self.output_dir / f"{stem}_adapted_{timestamp}{input_path.suffix}"

        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Загрузить и обработать
        if self.mode == "com" and input_path.suffix.lower() in (".docx", ".doc"):
            self._adapt_with_com(input_path, template, output_path, rules)
        elif self.mode == "native" and input_path.suffix.lower() == ".docx":
            self._adapt_with_native(input_path, template, output_path, rules)
        else:
            # Fallback: просто копируем с применением текстовых правил
            self._adapt_fallback(input_path, output_path, rules)

        logger.info(f"Документ адаптирован: {output_path}")
        return str(output_path)

    def compare_documents(self, old_path: Optional[str], new_path: Optional[str]) -> str:
        """
        Сравнить два документа и вернуть HTML-отчёт с различиями.

        Возвращает путь к HTML-файлу с отчётом.
        """
        if not old_path or not new_path:
            raise AdaptationError("Для сравнения нужны оба пути: old_path и new_path")

        old_text = self._extract_text(old_path)
        new_text = self._extract_text(new_path)

        # Генерация HTML-диффа
        diff = difflib.HtmlDiff(tabsize=4, wrapcolumn=80)
        html = diff.make_file(
            old_text.splitlines(),
            new_text.splitlines(),
            fromdesc=f"Оригинал: {Path(old_path).name}",
            todesc=f"Новая версия: {Path(new_path).name}",
            charset="utf-8"
        )

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        report_path = self.output_dir / f"compare_report_{timestamp}.html"
        report_path.parent.mkdir(parents=True, exist_ok=True)
        report_path.write_text(html, encoding="utf-8")

        logger.info(f"Отчёт сравнения сохранён: {report_path}")
        return str(report_path)

    def batch_replace(self, doc_path: Optional[str], replacements: Dict[str, str]) -> str:
        """
        Массовая замена текста в документе.

        replacements: {"СТАРЫЙ_ТЕКСТ": "НОВЫЙ_ТЕКСТ", ...}

        Возвращает путь к изменённому файлу.
        """
        if not doc_path:
            raise AdaptationError("Не указан doc_path")
        if not replacements:
            raise AdaptationError("Словарь replacements пуст")

        path = Path(doc_path)
        if not path.exists():
            raise AdaptationError(error_file_not_found(doc_path))

        # Создаём копию для редактирования
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.output_dir / f"{path.stem}_replaced_{timestamp}{path.suffix}"
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if self.mode == "com" and path.suffix.lower() in (".docx", ".doc"):
            self._batch_replace_com(path, replacements, output_path)
        elif self.mode == "native" and path.suffix.lower() == ".docx":
            self._batch_replace_native(path, replacements, output_path)
        else:
            # Fallback: простая замена в тексте
            text = path.read_text(encoding="utf-8", errors="ignore")
            for old, new in replacements.items():
                text = text.replace(old, new)
            output_path.write_text(text, encoding="utf-8")

        logger.info(f"Массовая замена завершена: {output_path}")
        return str(output_path)

    def update_references(self, doc_path: Optional[str]) -> str:
        """
        Обновить нормативные ссылки в документе.

        Заменяет устаревшие ГОСТы, ФГОСы и другие нормативные документы
        на актуальные версии.

        Возвращает путь к обновлённому файлу.
        """
        if not doc_path:
            raise AdaptationError("Не указан doc_path")

        path = Path(doc_path)
        if not path.exists():
            raise AdaptationError(error_file_not_found(doc_path))

        # Справочник устаревших → актуальных ссылок
        reference_map = {
            # ФГОС ВО
            "ФГОС ВО 3++": "ФГОС ВО 3++ (ред. 2023)",
            "ФГОС ВО 3+": "ФГОС ВО 3++ (ред. 2023)",
            # ФГОС СПО
            "ФГОС СПО 3++": "ФГОС СПО 3++ (ред. 2023)",
            # ГОСТы
            "ГОСТ Р 7.0.11-2011": "ГОСТ Р 7.0.97-2016",
            "ГОСТ 7.32-2001": "ГОСТ Р 7.0.97-2016",
            # Приказы
            "Приказ Минобрнауки № 636": "Приказ Минобрнауки России № 636 (ред. 2024)",
            "Приказ Минобрнауки от 12.08.2020 № 1036": "Приказ Минобрнауки России от 12.08.2020 № 1036 (ред. 2024)",
        }

        # Дополняем пользовательскими правилами из конфига (если есть)
        custom_refs = getattr(self.config, "reference_updates", {})
        if isinstance(custom_refs, dict):
            reference_map.update(custom_refs)

        return self.batch_replace(str(path), reference_map)

    def apply_formatting_rules(self, doc_path: Optional[str], rules: Dict[str, Any]) -> str:
        """
        Применить правила форматирования к документу.

        rules может содержать:
            - font_name: str
            - font_size: int
            - line_spacing: float
            - margins: {"top": float, "bottom": float, "left": float, "right": float} (в см)
            - alignment: str (left, center, right, justify)
            - heading_style: dict с параметрами заголовков

        Возвращает путь к отформатированному файлу.
        """
        if not doc_path:
            raise AdaptationError("Не указан doc_path")

        path = Path(doc_path)
        if not path.exists():
            raise AdaptationError(error_file_not_found(doc_path))

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        output_path = self.output_dir / f"{path.stem}_formatted_{timestamp}{path.suffix}"
        output_path.parent.mkdir(parents=True, exist_ok=True)

        if self.mode == "com" and path.suffix.lower() in (".docx", ".doc"):
            self._apply_formatting_com(path, rules, output_path)
        elif self.mode == "native" and path.suffix.lower() == ".docx":
            self._apply_formatting_native(path, rules, output_path)
        else:
            logger.warning("Форматирование для данного типа файлов не поддерживается — копирую без изменений")
            import shutil
            shutil.copy2(path, output_path)

        logger.info(f"Форматирование применено: {output_path}")
        return str(output_path)

    # ------------------------------------------------------------------
    # Internal helpers — text extraction
    # ------------------------------------------------------------------

    def _extract_text(self, file_path: str) -> str:
        """Извлечь текст из документа любого поддерживаемого формата."""
        path = Path(file_path)
        if not path.exists():
            return ""

        suffix = path.suffix.lower()

        if suffix == ".txt":
            return path.read_text(encoding="utf-8", errors="ignore")

        if suffix == ".docx" and _HAS_DOCX:
            try:
                doc = Document(str(path))
                paragraphs = [p.text for p in doc.paragraphs]
                # Таблицы
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            paragraphs.append(cell.text)
                return "\n".join(paragraphs)
            except Exception as e:
                logger.warning(f"Не удалось извлечь текст из {path}: {e}")

        if suffix in (".doc", ".docx") and _HAS_WIN32:
            try:
                word = win32.Dispatch("Word.Application")
                word.Visible = False
                doc = word.Documents.Open(str(path.absolute()))
                text = doc.Content.Text
                doc.Close(SaveChanges=False)
                word.Quit()
                return text
            except Exception as e:
                logger.warning(f"COM-извлечение текста не удалось: {e}")

        # Последняя попытка — как текст
        try:
            return path.read_text(encoding="utf-8", errors="ignore")
        except Exception:
            return ""

    # ------------------------------------------------------------------
    # COM implementations (Windows + Office)
    # ------------------------------------------------------------------

    def _adapt_with_com(self, input_path: Path, template: Optional[str],
                        output_path: Path, rules: Dict[str, Any]) -> None:
        """Адаптация через Word COM."""
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(input_path.absolute()))

            # Применить шаблон стилей, если указан
            if template:
                template_path = Path(template)
                if not template_path.exists():
                    template_path = self.templates_dir / template
                if template_path.exists():
                    doc.AttachedTemplate = str(template_path.absolute())

            # Применить текстовые правила
            for old_text, new_text in rules.get("replacements", {}).items():
                self._word_find_replace(doc, old_text, new_text)

            # Применить форматирование
            fmt = rules.get("formatting", {})
            if fmt:
                self._apply_word_formatting(doc, fmt)

            doc.SaveAs2(str(output_path.absolute()))
            doc.Close(SaveChanges=False)
        finally:
            word.Quit()

    def _batch_replace_com(self, path: Path, replacements: Dict[str, str],
                           output_path: Path) -> None:
        """Массовая замена через Word COM."""
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(path.absolute()))
            for old_text, new_text in replacements.items():
                self._word_find_replace(doc, old_text, new_text)
            doc.SaveAs2(str(output_path.absolute()))
            doc.Close(SaveChanges=False)
        finally:
            word.Quit()

    def _apply_formatting_com(self, path: Path, rules: Dict[str, Any],
                              output_path: Path) -> None:
        """Применить форматирование через Word COM."""
        word = win32.Dispatch("Word.Application")
        word.Visible = False
        try:
            doc = word.Documents.Open(str(path.absolute()))
            self._apply_word_formatting(doc, rules)
            doc.SaveAs2(str(output_path.absolute()))
            doc.Close(SaveChanges=False)
        finally:
            word.Quit()

    def _word_find_replace(self, doc, old_text: str, new_text: str) -> None:
        """Найти и заменить в Word COM."""
        find = doc.Content.Find
        find.ClearFormatting()
        find.Replacement.ClearFormatting()
        find.Text = old_text
        find.Replacement.Text = new_text
        find.Forward = True
        find.Wrap = 1  # wdFindContinue
        find.Format = False
        find.MatchCase = False
        find.MatchWholeWord = False
        find.Execute(Replace=2)  # wdReplaceAll

    def _apply_word_formatting(self, doc, rules: Dict[str, Any]) -> None:
        """Применить форматирование к документу Word COM."""
        # Поля документа
        if "margins" in rules:
            m = rules["margins"]
            page_setup = doc.PageSetup
            if "top" in m:
                page_setup.TopMargin = self._cm_to_points(m["top"])
            if "bottom" in m:
                page_setup.BottomMargin = self._cm_to_points(m["bottom"])
            if "left" in m:
                page_setup.LeftMargin = self._cm_to_points(m["left"])
            if "right" in m:
                page_setup.RightMargin = self._cm_to_points(m["right"])

        # Стили абзацев
        font_name = rules.get("font_name")
        font_size = rules.get("font_size")
        line_spacing = rules.get("line_spacing")
        alignment_map = {
            "left": 0,      # wdAlignParagraphLeft
            "center": 1,    # wdAlignParagraphCenter
            "right": 2,     # wdAlignParagraphRight
            "justify": 3,   # wdAlignParagraphJustify
        }
        alignment = rules.get("alignment")

        for paragraph in doc.Paragraphs:
            if font_name:
                paragraph.Range.Font.Name = font_name
            if font_size:
                paragraph.Range.Font.Size = font_size
            if line_spacing:
                paragraph.LineSpacingRule = 0  # wdLineSpaceMultiple (приближённо)
                paragraph.LineSpacing = line_spacing * 12  # приблизительно в pt
            if alignment and alignment in alignment_map:
                paragraph.Alignment = alignment_map[alignment]

        # Заголовки
        heading_style = rules.get("heading_style", {})
        if heading_style:
            for paragraph in doc.Paragraphs:
                if paragraph.Style.NameLocal.startswith("Заголовок") or \
                   paragraph.Style.NameLocal.startswith("Heading"):
                    if "font_name" in heading_style:
                        paragraph.Range.Font.Name = heading_style["font_name"]
                    if "font_size" in heading_style:
                        paragraph.Range.Font.Size = heading_style["font_size"]
                    if "bold" in heading_style:
                        paragraph.Range.Font.Bold = heading_style["bold"]

    def _cm_to_points(self, cm: float) -> float:
        """Перевести сантиметры в пункты Word."""
        return cm * 28.3465

    # ------------------------------------------------------------------
    # Native implementations (python-docx)
    # ------------------------------------------------------------------

    def _adapt_with_native(self, input_path: Path, template: Optional[str],
                           output_path: Path, rules: Dict[str, Any]) -> None:
        """Адаптация через python-docx."""
        doc = Document(str(input_path))

        # Применить замены
        for old_text, new_text in rules.get("replacements", {}).items():
            self._docx_replace(doc, old_text, new_text)

        # Применить форматирование
        fmt = rules.get("formatting", {})
        if fmt:
            self._apply_docx_formatting(doc, fmt)

        doc.save(str(output_path))

    def _batch_replace_native(self, path: Path, replacements: Dict[str, str],
                              output_path: Path) -> None:
        """Массовая замена через python-docx."""
        doc = Document(str(path))
        for old_text, new_text in replacements.items():
            self._docx_replace(doc, old_text, new_text)
        doc.save(str(output_path))

    def _apply_formatting_native(self, path: Path, rules: Dict[str, Any],
                                 output_path: Path) -> None:
        """Применить форматирование через python-docx."""
        doc = Document(str(path))
        self._apply_docx_formatting(doc, rules)
        doc.save(str(output_path))

    def _docx_replace(self, doc, old_text: str, new_text: str) -> None:
        """Найти и заменить текст в python-docx документе."""
        # Замена в параграфах
        for paragraph in doc.paragraphs:
            if old_text in paragraph.text:
                self._replace_in_paragraph(paragraph, old_text, new_text)

        # Замена в таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if old_text in paragraph.text:
                            self._replace_in_paragraph(paragraph, old_text, new_text)

    def _replace_in_paragraph(self, paragraph, old_text: str, new_text: str) -> None:
        """Заменить текст в параграфе python-docx, сохраняя форматирование ранов."""
        if old_text not in paragraph.text:
            return

        # Если весь параграф — один run, просто заменяем
        if len(paragraph.runs) == 1:
            paragraph.runs[0].text = paragraph.runs[0].text.replace(old_text, new_text)
            return

        # Собираем полный текст и индексы
        full_text = ""
        run_map = []
        for i, run in enumerate(paragraph.runs):
            run_text = run.text
            run_map.append((len(full_text), len(full_text) + len(run_text), run))
            full_text += run_text

        # Находим все вхождения
        start = 0
        occurrences = []
        while True:
            idx = full_text.find(old_text, start)
            if idx == -1:
                break
            occurrences.append(idx)
            start = idx + 1

        if not occurrences:
            return

        # Заменяем справа налево, чтобы индексы не съезжали
        for idx in reversed(occurrences):
            end_idx = idx + len(old_text)
            # Находим run, содержащий начало вхождения
            for start_pos, end_pos, run in run_map:
                if start_pos <= idx < end_pos:
                    # Если вхождение целиком внутри одного run
                    if end_idx <= end_pos:
                        local_start = idx - start_pos
                        local_end = end_idx - start_pos
                        run.text = run.text[:local_start] + new_text + run.text[local_end:]
                    else:
                        # Вхождение разбито на несколько run — заменяем в первом run целиком
                        local_start = idx - start_pos
                        run.text = run.text[:local_start] + new_text
                        # Очищаем остальные run, попавшие в диапазон
                        for s2, e2, r2 in run_map:
                            if s2 >= idx and e2 <= end_idx and r2 != run:
                                r2.text = ""
                            elif s2 < end_idx < e2:
                                # Конец вхождения внутри этого run
                                overlap = end_idx - s2
                                r2.text = r2.text[overlap:]
                    break

    def _apply_docx_formatting(self, doc, rules: Dict[str, Any]) -> None:
        """Применить форматирование к документу python-docx."""
        # Поля документа
        if "margins" in rules:
            m = rules["margins"]
            sections = doc.sections
            for section in sections:
                if "top" in m:
                    section.top_margin = Inches(m["top"] / 2.54)
                if "bottom" in m:
                    section.bottom_margin = Inches(m["bottom"] / 2.54)
                if "left" in m:
                    section.left_margin = Inches(m["left"] / 2.54)
                if "right" in m:
                    section.right_margin = Inches(m["right"] / 2.54)

        font_name = rules.get("font_name")
        font_size = rules.get("font_size")
        line_spacing = rules.get("line_spacing")
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        alignment = rules.get("alignment")

        # Применить к стилю Normal
        style = doc.styles["Normal"]
        if font_name:
            style.font.name = font_name
        if font_size:
            style.font.size = Pt(font_size)

        # Применить к существующим параграфам
        for paragraph in doc.paragraphs:
            if alignment and alignment in alignment_map:
                paragraph.alignment = alignment_map[alignment]
            if line_spacing:
                paragraph.paragraph_format.line_spacing = line_spacing
            for run in paragraph.runs:
                if font_name:
                    run.font.name = font_name
                if font_size:
                    run.font.size = Pt(font_size)

        # Таблицы
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if alignment and alignment in alignment_map:
                            paragraph.alignment = alignment_map[alignment]
                        if line_spacing:
                            paragraph.paragraph_format.line_spacing = line_spacing
                        for run in paragraph.runs:
                            if font_name:
                                run.font.name = font_name
                            if font_size:
                                run.font.size = Pt(font_size)

        # Заголовки
        heading_style = rules.get("heading_style", {})
        if heading_style:
            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading") or \
                   paragraph.style.name.startswith("Заголовок"):
                    for run in paragraph.runs:
                        if "font_name" in heading_style:
                            run.font.name = heading_style["font_name"]
                        if "font_size" in heading_style:
                            run.font.size = Pt(heading_style["font_size"])
                        if "bold" in heading_style:
                            run.font.bold = heading_style["bold"]

    # ------------------------------------------------------------------
    # Fallback
    # ------------------------------------------------------------------

    def _adapt_fallback(self, input_path: Path, output_path: Path, rules: Dict[str, Any]) -> None:
        """Fallback: копирование с текстовой заменой."""
        import shutil
        shutil.copy2(input_path, output_path)

        # Если это текстовый файл, применяем замены
        if input_path.suffix.lower() in (".txt", ".md", ".html", ".xml"):
            text = output_path.read_text(encoding="utf-8", errors="ignore")
            for old, new in rules.get("replacements", {}).items():
                text = text.replace(old, new)
            output_path.write_text(text, encoding="utf-8")

        logger.info(f"Fallback адаптация (копирование): {output_path}")
