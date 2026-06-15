"""
Утилиты для работы с документами методиста.
"""

import os
import re
from pathlib import Path
from typing import Optional


def ensure_directory(path: str) -> Path:
    """Создаёт директорию, если она не существует."""
    p = Path(path).expanduser()
    p.mkdir(parents=True, exist_ok=True)
    return p


def generate_filename(template_name: str, subject: str, extension: str = "docx") -> str:
    """Генерирует имя файла на основе шаблона и темы."""
    safe_subject = re.sub(r'[\\/:*?"<>|]', "_", subject).strip()
    safe_template = re.sub(r'[\\/:*?"<>|]', "_", template_name).strip()
    if not safe_subject:
        safe_subject = "документ"
    return f"{safe_template}_{safe_subject}.{extension}"


def find_replace_docx(doc, find: str, replace: str) -> int:
    """
    Поиск и замена текста в docx с сохранением стилей.
    Работает с python-docx (native-режим).
    Возвращает количество замен.
    """
    count = 0
    find_lower = find.lower()

    # Замена в параграфах
    for paragraph in doc.paragraphs:
        if find_lower in paragraph.text.lower():
            for run in paragraph.runs:
                if find_lower in run.text.lower():
                    run.text = re.sub(re.escape(find), replace, run.text, flags=re.IGNORECASE)
                    count += 1

    # Замена в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if find_lower in paragraph.text.lower():
                        for run in paragraph.runs:
                            if find_lower in run.text.lower():
                                run.text = re.sub(re.escape(find), replace, run.text, flags=re.IGNORECASE)
                                count += 1

    # Замена в колонтитулах (headers/footers)
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header is not None:
                for paragraph in header.paragraphs:
                    if find_lower in paragraph.text.lower():
                        for run in paragraph.runs:
                            if find_lower in run.text.lower():
                                run.text = re.sub(re.escape(find), replace, run.text, flags=re.IGNORECASE)
                                count += 1
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for paragraph in footer.paragraphs:
                    if find_lower in paragraph.text.lower():
                        for run in paragraph.runs:
                            if find_lower in run.text.lower():
                                run.text = re.sub(re.escape(find), replace, run.text, flags=re.IGNORECASE)
                                count += 1

    return count


def apply_template(doc, template_path: str) -> None:
    """
    Применяет стили из шаблона к документу (native-режим, python-docx).
    Копирует стили из шаблона, если они отсутствуют в целевом документе.
    """
    from docx import Document

    tpl = Document(template_path)
    tpl_styles = {s.name: s for s in tpl.styles}
    doc_styles = {s.name: s for s in doc.styles}

    for name, tpl_style in tpl_styles.items():
        if name not in doc_styles:
            try:
                new_style = doc.styles.add_style(name, tpl_style.type)
                new_style.font.name = tpl_style.font.name
                new_style.font.size = tpl_style.font.size
                new_style.font.bold = tpl_style.font.bold
                new_style.font.italic = tpl_style.font.italic
                new_style.font.color.rgb = tpl_style.font.color.rgb
                if hasattr(tpl_style, "paragraph_format"):
                    pf = tpl_style.paragraph_format
                    npf = new_style.paragraph_format
                    npf.alignment = pf.alignment
                    npf.space_before = pf.space_before
                    npf.space_after = pf.space_after
                    npf.line_spacing = pf.line_spacing
            except Exception:
                # Стиль может уже существовать или быть встроенным — игнорируем
                pass
