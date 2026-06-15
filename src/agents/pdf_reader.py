"""
Агент чтения PDF для Методист-Агента.

Извлекает текст, таблицы, изображения и метаданные из PDF.
Поддерживает OCR для сканированных документов и конвертацию в DOCX.
Реализован graceful degradation: если какая-либо функция недоступна,
агент продолжает работу и сообщает об ограничениях.

Windows-first, русский язык.
"""

import os
import shutil
import tempfile
import warnings
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd
import pdfplumber
import pytesseract
from docx import Document
from docx.shared import Inches
from pdf2image import convert_from_path
from PyPDF2 import PdfReader

from core.config import Config, get_output_dir

# Windows: попытка найти Tesseract, если не в PATH
if os.name == "nt":
    _tesseract_paths = [
        r"C:\Program Files\Tesseract-OCR\tesseract.exe",
        r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe",
    ]
    for _tp in _tesseract_paths:
        if Path(_tp).exists():
            pytesseract.pytesseract.tesseract_cmd = _tp
            break


class PDFReaderAgent:
    """Агент для работы с PDF-документами."""

    def __init__(self, config: Config):
        self.config = config
        self.output_dir = get_output_dir(config)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Диспетчер
    # ------------------------------------------------------------------
    def execute(self, task_type: Any, params: Dict[str, Any]) -> Any:
        """Диспетчер задач агента."""
        task_str = str(task_type).lower()

        pdf_path = params.get("pdf_path") or params.get("input_file", "")
        if not pdf_path:
            return {"error": "Не указан путь к PDF-файлу"}

        if "read" in task_str or "extract" in task_str or "текст" in task_str:
            return self.extract_text(pdf_path)

        if "table" in task_str or "таблиц" in task_str:
            return self.extract_tables(pdf_path)

        if "image" in task_str or "изображен" in task_str:
            output_dir = params.get("output_dir", "")
            return self.extract_images(pdf_path, output_dir)

        if "ocr" in task_str or "скан" in task_str:
            return self.ocr_pdf(pdf_path)

        if "convert" in task_str or "конверт" in task_str or "docx" in task_str:
            output_path = params.get("output_path", "")
            return self.convert_to_docx(pdf_path, output_path)

        if "metadata" in task_str or "метаданн" in task_str:
            return self.get_metadata(pdf_path)

        # По умолчанию — универсальная обработка
        return self.process(params)

    # ------------------------------------------------------------------
    # Извлечение текста
    # ------------------------------------------------------------------
    def extract_text(self, pdf_path: str) -> Dict[str, Any]:
        """Извлечение текста из PDF через pdfplumber."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        try:
            full_text = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = page.extract_text()
                    if text:
                        full_text.append(f"--- Страница {i} ---\n{text}")

            result = "\n\n".join(full_text)
            return {
                "pdf_path": str(path),
                "text": result,
                "pages": len(full_text),
                "success": True,
            }
        except Exception as e:
            return {"error": f"Ошибка извлечения текста: {e}", "pdf_path": str(path)}

    # ------------------------------------------------------------------
    # Извлечение таблиц
    # ------------------------------------------------------------------
    def extract_tables(self, pdf_path: str) -> Dict[str, Any]:
        """Извлечение таблиц из PDF в список DataFrame."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        tables_data: List[Dict[str, Any]] = []
        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    tables = page.extract_tables()
                    for j, table in enumerate(tables, start=1):
                        if table and len(table) > 1:
                            # Первая строка — заголовки, остальные — данные
                            headers = table[0]
                            rows = table[1:]
                            df = pd.DataFrame(rows, columns=headers)
                            tables_data.append({
                                "page": i,
                                "table_index": j,
                                "columns": list(df.columns),
                                "rows_count": len(df),
                                "data": df.to_dict(orient="records"),
                            })

            return {
                "pdf_path": str(path),
                "tables_count": len(tables_data),
                "tables": tables_data,
                "success": True,
            }
        except Exception as e:
            return {"error": f"Ошибка извлечения таблиц: {e}", "pdf_path": str(path)}

    # ------------------------------------------------------------------
    # Извлечение изображений
    # ------------------------------------------------------------------
    def extract_images(self, pdf_path: str, output_dir: Optional[str] = None) -> Dict[str, Any]:
        """Извлечение изображений из PDF через pdf2image."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        if output_dir:
            out = Path(output_dir)
        else:
            out = self.output_dir / "pdf_images" / path.stem
        out.mkdir(parents=True, exist_ok=True)

        try:
            images = convert_from_path(str(path), dpi=200)
            saved_paths: List[str] = []
            for i, image in enumerate(images, start=1):
                img_path = out / f"page_{i:03d}.png"
                image.save(str(img_path), "PNG")
                saved_paths.append(str(img_path))

            return {
                "pdf_path": str(path),
                "output_dir": str(out),
                "images_count": len(saved_paths),
                "images": saved_paths,
                "success": True,
            }
        except Exception as e:
            # Graceful degradation: если pdf2image/poppler недоступен
            return {
                "error": f"Не удалось извлечь изображения (возможно, не установлен poppler): {e}",
                "pdf_path": str(path),
                "hint": "Установите poppler для Windows: https://github.com/oschwartz10612/poppler-windows",
            }

    # ------------------------------------------------------------------
    # OCR
    # ------------------------------------------------------------------
    def ocr_pdf(self, pdf_path: str) -> Dict[str, Any]:
        """OCR для сканированных PDF: конвертирует страницы в изображения и распознаёт текст."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        # Проверяем доступность Tesseract
        try:
            pytesseract.get_tesseract_version()
        except Exception as e:
            return {
                "error": f"Tesseract OCR недоступен: {e}",
                "hint": "Установите Tesseract: https://github.com/UB-Mannheim/tesseract/wiki",
            }

        try:
            images = convert_from_path(str(path), dpi=300)
            ocr_texts = []
            for i, image in enumerate(images, start=1):
                text = pytesseract.image_to_string(image, lang="rus+eng")
                if text.strip():
                    ocr_texts.append(f"--- Страница {i} ---\n{text.strip()}")

            return {
                "pdf_path": str(path),
                "text": "\n\n".join(ocr_texts),
                "pages": len(ocr_texts),
                "success": True,
            }
        except Exception as e:
            return {
                "error": f"Ошибка OCR: {e}",
                "pdf_path": str(path),
                "hint": "Убедитесь, что установлены poppler и Tesseract с русским языковым пакетом",
            }

    # ------------------------------------------------------------------
    # Конвертация в DOCX
    # ------------------------------------------------------------------
    def convert_to_docx(self, pdf_path: str, output_path: Optional[str] = None) -> Dict[str, Any]:
        """Конвертация PDF в DOCX: текст + таблицы + изображения (если возможно)."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        if not output_path:
            output_path = str(self.output_dir / f"{path.stem}.docx")
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)

        doc = Document()
        warnings_list: List[str] = []

        try:
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    # Текст страницы
                    text = page.extract_text()
                    if text:
                        for line in text.splitlines():
                            if line.strip():
                                doc.add_paragraph(line.strip())

                    # Таблицы страницы
                    tables = page.extract_tables()
                    for table in tables:
                        if table and len(table) > 1:
                            tbl = doc.add_table(rows=len(table), cols=len(table[0]))
                            tbl.style = "Table Grid"
                            for r_idx, row in enumerate(table):
                                for c_idx, cell in enumerate(row):
                                    tbl.rows[r_idx].cells[c_idx].text = str(cell or "")
                            doc.add_paragraph()  # отступ после таблицы

            # Попытка вставить изображения (graceful degradation)
            try:
                images = convert_from_path(str(path), dpi=150)
                for i, image in enumerate(images, start=1):
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        image.save(tmp.name, "PNG")
                        tmp_path = tmp.name
                    try:
                        doc.add_picture(tmp_path, width=Inches(6.0))
                        doc.add_paragraph()
                    except Exception:
                        warnings_list.append(f"Не удалось вставить изображение страницы {i}")
                    finally:
                        Path(tmp_path).unlink(missing_ok=True)
            except Exception as e:
                warnings_list.append(f"Изображения не извлечены: {e}")

            doc.save(str(out))
            return {
                "pdf_path": str(path),
                "output_path": str(out),
                "success": True,
                "warnings": warnings_list,
            }
        except Exception as e:
            return {"error": f"Ошибка конвертации в DOCX: {e}", "pdf_path": str(path)}

    # ------------------------------------------------------------------
    # Метаданные
    # ------------------------------------------------------------------
    def get_metadata(self, pdf_path: str) -> Dict[str, Any]:
        """Получение метаданных PDF через PyPDF2."""
        path = Path(pdf_path)
        if not path.exists():
            return {"error": f"Файл не найден: {pdf_path}"}

        try:
            reader = PdfReader(str(path))
            meta = reader.metadata or {}
            info = {
                "pdf_path": str(path),
                "pages": len(reader.pages),
                "title": meta.get("/Title", ""),
                "author": meta.get("/Author", ""),
                "subject": meta.get("/Subject", ""),
                "creator": meta.get("/Creator", ""),
                "producer": meta.get("/Producer", ""),
                "creation_date": meta.get("/CreationDate", ""),
                "modification_date": meta.get("/ModDate", ""),
                "success": True,
            }
            return info
        except Exception as e:
            return {"error": f"Ошибка чтения метаданных: {e}", "pdf_path": str(path)}

    # ------------------------------------------------------------------
    # Универсальный метод
    # ------------------------------------------------------------------
    def process(self, params: Dict[str, Any]) -> Dict[str, Any]:
        """Универсальная обработка PDF: извлекает всё, что возможно."""
        pdf_path = params.get("pdf_path") or params.get("input_file", "")
        if not pdf_path:
            return {"error": "Не указан путь к PDF-файлу"}

        result: Dict[str, Any] = {
            "pdf_path": pdf_path,
            "text": None,
            "tables": None,
            "metadata": None,
            "ocr": None,
            "success": False,
        }

        # 1. Метаданные
        meta = self.get_metadata(pdf_path)
        if meta.get("success"):
            result["metadata"] = meta
        else:
            result["metadata_error"] = meta.get("error")

        # 2. Текст
        text_res = self.extract_text(pdf_path)
        if text_res.get("success"):
            result["text"] = text_res
        else:
            result["text_error"] = text_res.get("error")

        # 3. Таблицы
        tables_res = self.extract_tables(pdf_path)
        if tables_res.get("success"):
            result["tables"] = tables_res
        else:
            result["tables_error"] = tables_res.get("error")

        # 4. OCR (если текст пустой или явно запрощено)
        action = params.get("action", "")
        if action == "ocr" or (text_res.get("success") and not text_res.get("text", "").strip()):
            ocr_res = self.ocr_pdf(pdf_path)
            if ocr_res.get("success"):
                result["ocr"] = ocr_res
            else:
                result["ocr_error"] = ocr_res.get("error")

        # 5. Конвертация (если запрошено)
        if action == "convert" or "docx" in action:
            output = params.get("output_path", "")
            conv_res = self.convert_to_docx(pdf_path, output)
            if conv_res.get("success"):
                result["conversion"] = conv_res
            else:
                result["conversion_error"] = conv_res.get("error")

        result["success"] = any(
            result.get(k) and result[k].get("success")
            for k in ["text", "tables", "metadata", "ocr", "conversion"]
        )
        return result
